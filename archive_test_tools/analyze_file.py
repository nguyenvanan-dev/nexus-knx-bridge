#!/usr/bin/env python3
from pathlib import Path
import sys
import json

def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

def read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"\n--- PAGE {i} ---\n{text}")
    return "\n".join(parts)

def read_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def read_xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"\n--- SHEET: {ws.title} ---")
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in vals):
                parts.append(" | ".join(vals))
    return "\n".join(parts)

def detect_kind(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in [".xlsx", ".xlsm"]:
        return "xlsx"
    if ext in [".txt", ".md", ".csv"]:
        return "txt"
    return "unknown"

def main():
    if len(sys.argv) < 2:
        print("Usage: analyze_file.py <file_path>")
        sys.exit(1)

    path = Path(sys.argv[1]).expanduser().resolve()
    if not path.exists():
        print(json.dumps({"ok": False, "error": "file_not_found", "path": str(path)}, ensure_ascii=False))
        sys.exit(1)

    kind = detect_kind(path)

    if kind == "pdf":
        text = read_pdf(path)
    elif kind == "docx":
        text = read_docx(path)
    elif kind == "xlsx":
        text = read_xlsx(path)
    elif kind == "txt":
        text = read_txt(path)
    else:
        print(json.dumps({"ok": False, "error": "unsupported_file_type", "file": path.name}, ensure_ascii=False))
        sys.exit(1)

    extracted_dir = Path.home() / ".openclaw/workspace/knowledge/extracted"
    extracted_dir.mkdir(parents=True, exist_ok=True)

    out_path = extracted_dir / f"{path.stem}.txt"
    out_path.write_text(text, encoding="utf-8")

    preview = text[:3000]

    result = {
        "ok": True,
        "file": path.name,
        "kind": kind,
        "text_chars": len(text),
        "extracted_path": str(out_path),
        "preview": preview
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
