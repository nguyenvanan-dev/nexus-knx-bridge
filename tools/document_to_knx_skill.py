#!/usr/bin/env python3
import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse, parse_qs

import requests
import pandas as pd
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document


WORKSPACE = Path.home() / ".openclaw" / "workspace"
KNOWLEDGE = WORKSPACE / "knowledge"

INBOX = KNOWLEDGE / "inbox"
EXTRACTED = KNOWLEDGE / "extracted"
SUMMARIES = KNOWLEDGE / "summaries"
REVIEW = KNOWLEDGE / "review"

DEVICE_JSON = Path.home() / "knx-bridge" / "devices.json"


def now_stamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def ensure_dirs():
    for p in [INBOX, EXTRACTED, SUMMARIES, REVIEW]:
        p.mkdir(parents=True, exist_ok=True)


def safe_name(name: str) -> str:
    name = name.strip().replace(" ", "_")
    name = re.sub(r"[^A-Za-z0-9_.-]+", "_", name)
    return name[:120] or f"source_{now_stamp()}"


def is_url(source: str) -> bool:
    return source.startswith("http://") or source.startswith("https://")


def google_sheet_export_url(url: str) -> str | None:
    """
    Chuyển link Google Sheet public thành link export xlsx.
    Ví dụ:
    https://docs.google.com/spreadsheets/d/<ID>/edit#gid=0
    """
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", url)
    if not m:
        return None
    sheet_id = m.group(1)
    return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx"


def filename_from_response(resp, fallback: str):
    cd = resp.headers.get("content-disposition", "")
    m = re.search(r'filename="?([^";]+)"?', cd)
    if m:
        return safe_name(m.group(1))

    path = urlparse(resp.url).path
    base = os.path.basename(path)
    if base and "." in base:
        return safe_name(base)

    ctype = resp.headers.get("content-type", "").lower()
    if "pdf" in ctype:
        return fallback + ".pdf"
    if "spreadsheet" in ctype or "excel" in ctype:
        return fallback + ".xlsx"
    if "word" in ctype:
        return fallback + ".docx"
    if "html" in ctype:
        return fallback + ".html"
    return fallback + ".bin"


def save_source(source: str) -> Path:
    ensure_dirs()

    stamp = now_stamp()

    if is_url(source):
        url = source.strip()

        export = google_sheet_export_url(url)
        if export:
            url = export

        headers = {
            "User-Agent": "KNX-SmartHome-Assistant/1.0"
        }

        try:
            resp = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
        except Exception as e:
            raise RuntimeError(f"Không tải được link: {e}")

        if resp.status_code >= 400:
            raise RuntimeError(f"Link trả về HTTP {resp.status_code}. Có thể link chưa public hoặc cần đăng nhập.")

        fallback = f"link_{stamp}"
        fname = filename_from_response(resp, fallback)
        dest = INBOX / fname

        dest.write_bytes(resp.content)

        meta = {
            "source_type": "url",
            "original_url": source,
            "downloaded_url": resp.url,
            "status_code": resp.status_code,
            "content_type": resp.headers.get("content-type"),
            "saved_as": str(dest),
            "received_at": datetime.now().isoformat(timespec="seconds")
        }
        (INBOX / f"{dest.stem}.meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
        return dest

    src = Path(source).expanduser().resolve()
    if not src.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {src}")

    dest = INBOX / f"{stamp}_{safe_name(src.name)}"
    shutil.copy2(src, dest)
    return dest


def read_txt(path: Path) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1"]:
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        parts.append(f"\n\n--- PAGE {i} ---\n{text}")
    return "\n".join(parts)


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for ti, table in enumerate(doc.tables, start=1):
        parts.append(f"\n--- TABLE {ti} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_excel(path: Path) -> str:
    xls = pd.ExcelFile(path)
    parts = []

    for sheet in xls.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet, dtype=str)
        df = df.fillna("")
        parts.append(f"\n--- SHEET: {sheet} ---")
        parts.append(df.to_csv(index=False))

    return "\n".join(parts)


def read_csv(path: Path) -> str:
    return read_txt(path)


def read_html(path: Path) -> str:
    html = read_txt(path)
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text("\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()

    if ext in [".txt", ".md", ".log"]:
        return read_txt(path)

    if ext == ".pdf":
        return read_pdf(path)

    if ext == ".docx":
        return read_docx(path)

    if ext in [".xlsx", ".xls"]:
        return read_excel(path)

    if ext == ".csv":
        return read_csv(path)

    if ext in [".html", ".htm"]:
        return read_html(path)

    raise RuntimeError(f"Chưa hỗ trợ định dạng file: {ext}")


def detect_candidates(text: str):
    group_addresses = sorted(set(re.findall(r"\b\d{1,2}/\d{1,2}/\d{1,3}\b", text)))

    dpts = sorted(set(re.findall(r"\b(?:DPT\s*)?(\d{1,3}\.\d{1,3})\b", text, flags=re.IGNORECASE)))

    keywords = {
        "light": ["đèn", "den", "light", "lighting", "lamp"],
        "dimmer": ["dimmer", "dim", "dimming"],
        "curtain": ["rèm", "rem", "curtain", "blind", "shutter"],
        "ac": ["điều hòa", "dieu hoa", "ac", "air conditioner", "hvac"],
        "sensor": ["sensor", "cảm biến", "cam bien"],
        "irrigation": ["tưới", "tuoi", "irrigation", "watering"],
        "audio": ["audio", "loa", "speaker", "music"]
    }

    found_types = []
    low = text.lower()
    for typ, keys in keywords.items():
        if any(k in low for k in keys):
            found_types.append(typ)

    return {
        "group_addresses": group_addresses,
        "dpts": dpts,
        "possible_device_types": found_types
    }


def load_existing_devices():
    if not DEVICE_JSON.exists():
        return None

    try:
        return json.loads(DEVICE_JSON.read_text(encoding="utf-8"))
    except Exception:
        return None


def make_summary(source_path: Path, text: str, candidates: dict) -> str:
    lines = []
    lines.append(f"# Summary for {source_path.name}")
    lines.append("")
    lines.append(f"- Source file: `{source_path}`")
    lines.append(f"- Extracted at: {datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"- Text length: {len(text)} characters")
    lines.append(f"- Group addresses found: {len(candidates['group_addresses'])}")
    lines.append(f"- DPT values found: {len(candidates['dpts'])}")
    lines.append(f"- Possible device types: {', '.join(candidates['possible_device_types']) or 'unknown'}")
    lines.append("")
    lines.append("## Group addresses")
    for ga in candidates["group_addresses"][:200]:
        lines.append(f"- {ga}")
    lines.append("")
    lines.append("## DPT")
    for dpt in candidates["dpts"][:100]:
        lines.append(f"- {dpt}")
    return "\n".join(lines)


def make_proposal(source_path: Path, extracted_path: Path, summary_path: Path, text: str, candidates: dict):
    existing = load_existing_devices()

    proposal = {
        "proposal_type": "document_to_knx_config",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "source": {
            "file": str(source_path),
            "extracted_text": str(extracted_path),
            "summary": str(summary_path)
        },
        "status": "needs_ai_review",
        "important_warning": [
            "Đây mới là proposal nháp.",
            "Không được ghi thẳng vào devices.json.",
            "Bot cần đọc extracted_text và tạo proposed_devices có cấu trúc chuẩn.",
            "Người dùng phải xác nhận trước khi cập nhật devices.json."
        ],
        "detected": candidates,
        "existing_devices_loaded": existing is not None,
        "proposed_devices": [],
        "missing_info": [],
        "next_instruction_for_bot": (
            "Hãy đọc file extracted_text, trích xuất từng thiết bị KNX. "
            "Với mỗi thiết bị cần có name, type, room, functions, group_address, dpt, direction. "
            "Nếu thiếu group_address hoặc dpt thì status=missing_info. "
            "Nếu đủ thì status=ready. "
            "Không tự đoán group_address."
        )
    }

    return proposal


def main():
    parser = argparse.ArgumentParser(description="Document to KNX Config Skill")
    parser.add_argument("source", help="Đường dẫn file local hoặc link public")
    args = parser.parse_args()

    ensure_dirs()

    source_path = save_source(args.source)

    text = extract_text(source_path)

    stem = source_path.stem
    extracted_path = EXTRACTED / f"{stem}.txt"
    extracted_path.write_text(text, encoding="utf-8")

    candidates = detect_candidates(text)

    summary = make_summary(source_path, text, candidates)
    summary_path = SUMMARIES / f"{stem}_summary.md"
    summary_path.write_text(summary, encoding="utf-8")

    proposal = make_proposal(source_path, extracted_path, summary_path, text, candidates)
    proposal_path = REVIEW / f"device_proposal_{now_stamp()}_{stem}.json"
    proposal_path.write_text(json.dumps(proposal, ensure_ascii=False, indent=2), encoding="utf-8")

    print("OK: Đã xử lý tài liệu")
    print(f"SOURCE:    {source_path}")
    print(f"EXTRACTED: {extracted_path}")
    print(f"SUMMARY:   {summary_path}")
    print(f"PROPOSAL:  {proposal_path}")
    print()
    print(f"Group addresses found: {len(candidates['group_addresses'])}")
    print(f"DPT values found:      {len(candidates['dpts'])}")
    print(f"Possible types:        {', '.join(candidates['possible_device_types']) or 'unknown'}")


if __name__ == "__main__":
    main()
